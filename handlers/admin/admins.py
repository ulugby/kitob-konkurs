from filters import IsAdmin,IsPrivate
from aiogram import types,html
from aiogram.filters import Command
from loader import dp,bot
from aiogram.types.reaction_type_emoji import ReactionTypeEmoji
from datetime import datetime, timedelta
import sqlite3
import random,os,json,uuid
from aiogram.fsm.context import FSMContext
from aiogram.fsm.state import State, StatesGroup
from aiogram.types import InlineKeyboardMarkup, InlineKeyboardButton
from aiogram.filters.callback_data import CallbackData
from aiogram.utils.keyboard import InlineKeyboardBuilder
from docx import Document

from aiogram.types import FSInputFile

DATABASE_FILE = "bot.db"


conn = sqlite3.connect('bot.db')
cursor = conn.cursor()

def db_connection():
    conn = sqlite3.connect(DATABASE_FILE)
    return conn

def create_channel_table():
    conn = db_connection()
    cursor = conn.cursor()
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS channels (
            id INTEGER PRIMARY KEY,
            username VARCHAR(30) NULL,
            name TEXT,
            telegram_id BIGINTEGER,
            users_count BIGINTEGER NULL,
            registration_date TEXT,
            invite_link TEXT NULL,
            invite_required BOOLEAN DEFAULT FALSE
        )
    ''')
    conn.commit()
    conn.close()

# Ensure the table is created at startup
create_channel_table()


def get_total_users():
    conn = sqlite3.connect(DATABASE_FILE)
    cursor = conn.cursor()
    cursor.execute("SELECT COUNT(*) FROM users")
    total_users = cursor.fetchone()[0]
    conn.close()
    return total_users

def get_today_users():
    conn = sqlite3.connect(DATABASE_FILE)
    cursor = conn.cursor()
    today = datetime.now().strftime("%Y-%m-%d")
    cursor.execute(f"SELECT COUNT(*) FROM users WHERE registration_date >= '{today}'")
    today_users = cursor.fetchone()[0]
    conn.close()
    return today_users

def get_yesterday_users():
    yesterday = (datetime.now() - timedelta(days=1)).strftime("%Y-%m-%d")
    conn = sqlite3.connect(DATABASE_FILE)
    cursor = conn.cursor()
    cursor.execute(f"SELECT COUNT(*) FROM users WHERE registration_date >= '{yesterday}' AND registration_date < '{datetime.now().strftime('%Y-%m-%d')}'")
    yesterday_users = cursor.fetchone()[0]
    conn.close()
    return yesterday_users

def get_month_users():
    first_day_of_month = datetime.now().replace(day=1).strftime("%Y-%m-%d")
    conn = sqlite3.connect(DATABASE_FILE)
    cursor = conn.cursor()
    cursor.execute(f"SELECT COUNT(*) FROM users WHERE registration_date >= '{first_day_of_month}'")
    month_users = cursor.fetchone()[0]
    conn.close()
    return month_users

def get_channel_by_id(channel_id):
    conn = sqlite3.connect('bot.db')
    cursor = conn.cursor()

    cursor.execute("SELECT * FROM channels WHERE telegram_id = ?", (channel_id,))
    channel_data = cursor.fetchone()
    
    conn.close()
    return channel_data

class AddChannel(StatesGroup):
    waiting_for_channel_id = State()
    waiting_for_invite_required = State()


@dp.message(Command('panel'), IsAdmin(), IsPrivate())
async def admin_panel(message: types.Message):
    inline_keyboard = InlineKeyboardBuilder()
    inline_keyboard.add(
        InlineKeyboardButton(
            text="📊Statistics",
            callback_data="statistics"
        ),
        InlineKeyboardButton(
            text="📢Kanallar ro'yxati",
            callback_data="list_channels"
        ),
        InlineKeyboardButton(
            text="❌",
            callback_data=f"deletemsg_{message.chat.id}"
        )
    )
    inline_keyboard.adjust(1)
    await message.answer("Assalamu alaykum admin panelga hush kelibsiz",reply_markup=inline_keyboard.as_markup())


@dp.callback_query(lambda query: query.data.startswith("back_panel"))
async def back_to_admin_panel(callback_query: types.CallbackQuery):
    inline_keyboard = InlineKeyboardBuilder()
    inline_keyboard.add(
        InlineKeyboardButton(
            text="📊Statistics",
            callback_data="statistics"
        ),
        InlineKeyboardButton(
            text="📢Kanallar ro'yxati",
            callback_data="list_channels"
        ),
        InlineKeyboardButton(
            text="❌",
            callback_data=f"deletemsg_{callback_query.message.chat.id}"
        )
    )
    inline_keyboard.adjust(1)
    await callback_query.message.edit_text("Assalamu alaykum admin panelga hush kelibsiz",reply_markup=inline_keyboard.as_markup())



@dp.callback_query(lambda query: query.data.startswith("deletemsg_"))
async def delmsg(callback_query: types.CallbackQuery):
    chat_id = int(callback_query.data.split('_')[-1])
    try:
        await bot.delete_message(chat_id=chat_id, message_id=callback_query.message.message_id)
    except:
        await callback_query.answer("Xatolik yuz berdi")


@dp.callback_query(lambda query: query.data.startswith("addchannel"))
async def add_channel_start(callback_query: types.CallbackQuery, state: FSMContext):
    bot_username = (await bot.me()).username    
    add_group_button = InlineKeyboardMarkup(
        inline_keyboard=[
            [
                InlineKeyboardButton(text="Kanalga admin qilish", url=f"https://t.me/{bot_username}?startchannel=true")
            ],
            [
                InlineKeyboardButton(text="Guruhga admin qilish", url=f"https://t.me/{bot_username}?startgroup=true")
            ],
            [
                InlineKeyboardButton(text="⏪Orqaga", callback_data="back_panel")
            ]
            
        ]
    )
    await callback_query.message.edit_text("<b>Kanal yoki guruhni ulash uchun ID ni yoki Xabarni yuboring\nKanal ulashdan avval botni Admin  ekanligini tekshiring</b>", reply_markup=add_group_button)
    await state.set_state(AddChannel.waiting_for_channel_id)


@dp.message(AddChannel.waiting_for_channel_id,IsAdmin(), IsPrivate())
async def add_channel_id(message: types.Message, state: FSMContext):
    inline_keyboard = InlineKeyboardBuilder()
    inline_keyboard.add(
        InlineKeyboardButton(
            text="⏪ Kanallar ro'yxatini ko'rish",
            callback_data=f"list_channels"
        )
    )
    if message.forward_from_chat:
        channel_id = message.forward_from_chat.id
    else:
        channel_id = message.text
    try:
        channel = await bot.get_chat(chat_id=channel_id)
        channel_username = channel.username
        channel_name = channel.title
        channel_users_count = await bot.get_chat_member_count(chat_id=channel_id)
        invite_link = await bot.export_chat_invite_link(chat_id=channel_id)

        # conn = db_connection()
        # cursor = conn.cursor()
        # cursor.execute('SELECT COUNT(*) FROM channels WHERE telegram_id = ?', (channel_id,))
        # if cursor.fetchone()[0] > 0:
        #     await message.answer(f"Bu {channel.type} allaqachon qo'shilgan. Boshqa kanal yoki guruh idsini tashlang", reply_markup=inline_keyboard.as_markup())
        #     await state.set_state(AddChannel.waiting_for_channel_id)
        #     conn.close()
        #     await state.clear()
        #     return
        
        # cursor.execute('''
        #     INSERT INTO channels (username, name, telegram_id, users_count, registration_date, invite_link)
        #     VALUES (?, ?, ?, ?, ?, ?)
        # ''', (channel_username, channel_name, channel_id, channel_users_count, datetime.now().strftime("%Y-%m-%d"), invite_link))
        # conn.commit()
        # conn.close()

        await state.update_data(
            channel_id=channel_id,
            username=channel_username,
            name=channel_name,
            users_count=channel_users_count,
            invite_link=invite_link,
        )
        inline_keyboard = InlineKeyboardBuilder()
        inline_keyboard.adjust(1)
        inline_keyboard.add(
            InlineKeyboardButton(text="✅ Ha", callback_data="invite_required_true"),
            InlineKeyboardButton(text="❌ Yo'q", callback_data="invite_required_false"),
        )

        await message.answer(
            "Kanalga qo'shilish so'rovli bo'lsinmi?\n\nTanlovingizni tugma orqali yuboring:",
            reply_markup=inline_keyboard.as_markup(),
        )
        await state.set_state(AddChannel.waiting_for_invite_required)
     
    except Exception as e:
        await message.answer(f"Kanal topilmadi yoki unga kirish imkoni yo'q: {e}")
        await state.clear()


@dp.callback_query(lambda c: c.data.startswith("invite_required_"),AddChannel.waiting_for_invite_required)
async def invite_required_callback(callback_query: types.CallbackQuery, state: FSMContext):
    # Callbackdan True yoki False qiymatini olish
    invite_required = callback_query.data == "invite_required_true"
    
    # Holatdan ma'lumotlarni olish
    data = await state.get_data()
    channel_id = data["channel_id"]
    username = data["username"]
    name = data["name"]
    users_count = data["users_count"]
    invite_link = data["invite_link"]

    if invite_required:
        try:
            invite_link = await bot.create_chat_invite_link(chat_id=channel_id, creates_join_request=True, name="Maxsus havola")
            invite_link = invite_link.invite_link
        except Exception as e:
            await callback_query.message.answer(f"Kanal uchun taklif havolasini yaratib bo'lmadi: {e}")
            await state.clear()
            return


    inline_keyboard = InlineKeyboardBuilder()
    inline_keyboard.adjust(1)
    inline_keyboard.add(
        InlineKeyboardButton(text="⏪ Kanallar ro'yxatini ko'rish", callback_data="list_channels"),
    )


    conn = db_connection()
    cursor = conn.cursor()
    cursor.execute('''
        INSERT INTO channels (username, name, telegram_id, users_count, registration_date, invite_link, invite_required)
        VALUES (?, ?, ?, ?, ?, ?, ?)
    ''', (username, name, channel_id, users_count, datetime.now().strftime("%Y-%m-%d"), invite_link, invite_required))
    conn.commit()
    conn.close()


    await callback_query.message.delete()
    await callback_query.message.answer(
        f"Kanal muvaffaqiyatli qo'shildi:\n\n"
        f"📌 Nomi: {name}\n"
        f"👥 A'zolar soni: {users_count}\n"
        f"✅ Qo'shilish talab: {'Ha' if invite_required else 'Yoq'}",
        reply_markup=inline_keyboard.as_markup()
    )
    await state.clear()

# @dp.message(Command('stat'))

@dp.callback_query(lambda query: query.data.startswith("statistics"))
async def statistic(callback_query: types.CallbackQuery):
    total_users = get_total_users()
    today_users = get_today_users()
    yesterday_users = get_yesterday_users()
    month_users = get_month_users()


    inline_keyboard = InlineKeyboardBuilder()
    inline_keyboard.add(
        InlineKeyboardButton(
            text="⏪Orqaga",
            callback_data="back_panel"
        ),
    )
    inline_keyboard.adjust(1)
    await callback_query.message.edit_text(
        f"📊 Bot Statistics\n\n"
        f"👤 Total members: {total_users}\n\n"
        f"📅 Members today: {today_users}\n"
        f"📅 Members yesterday: {yesterday_users}\n"
        f"📅 Members this month: {month_users}",
        reply_markup=inline_keyboard.as_markup()
        )

def get_channels():
    conn = sqlite3.connect('bot.db')
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM channels")
    channels = cursor.fetchall()
    conn.close()
    return channels



@dp.callback_query(lambda query: query.data.startswith("list_channels"))
async def list_channels(callback_query: types.CallbackQuery):
    channels = get_channels()
    inline_keyboard = InlineKeyboardBuilder()
    for channel in channels:
        channel_id = channel[3]
        channel_name = channel[2]
        inline_keyboard.add(
            InlineKeyboardButton(
                text=str(channel_name),
                callback_data=f"settings_{channel_id}"
            )
        )
    inline_keyboard.add(
        InlineKeyboardButton(
            text="➕Kanal qo'shish",
            callback_data="addchannel"
        ),
        InlineKeyboardButton(
            text="⏪Orqaga",
            callback_data="back_panel"
        ),
    )
    inline_keyboard.adjust(1)
    await callback_query.message.edit_text("Kanallar ro'yxati:", reply_markup=inline_keyboard.as_markup())

@dp.callback_query(lambda query: query.data.startswith("settings_"))
async def channel_settings_callback_handler(callback_query: types.CallbackQuery):
    channel_id = int(callback_query.data.split('_')[-1])
    inline_keyboard = InlineKeyboardBuilder()

    channel_data = get_channel_by_id(channel_id)
    print(channel_data)
    if channel_data:
        channel_name = channel_data[2]  # `name` ustuni
        users_count = channel_data[4]   # `users_count` ustuni
        username = channel_data[1]      # `username` ustuni
        invite_link = channel_data[6]   # `invite_link` ustuni

        if username:
            channel_url = f"https://t.me/{username}"
        else:
            channel_url = invite_link

        # Kanal ma'lumotlarini ko‘rsatish
        text = (
            f"<b>Kanal sozlamalari:</b>\n"
            f"📋 <b>ID:</b> {channel_id}\n"
            f"📌 <b>Nomi:</b> {channel_name}\n"
            f"👥 <b>A'zolar soni:</b> {users_count}\n"
        )
        inline_keyboard.add(
            InlineKeyboardButton(
                text="♻️Yangi link (invite link)",
                callback_data=f"new_invite_{channel_id}"
            ),
            InlineKeyboardButton(
                text="🗑️Kanalni o'chirish",
                callback_data=f"delete_{channel_id}"
            ),
            InlineKeyboardButton(
                text="🔗 Kanalga o'tish", 
                url=channel_url
            ),
            InlineKeyboardButton(
                text="⏪Orqaga",
                callback_data="list_channels"
            )
        )
        inline_keyboard.adjust(1)
        await callback_query.message.edit_text(text=text, reply_markup=inline_keyboard.as_markup())
    else:
        await callback_query.message.edit_text("❌Kanal topilmadi", reply_markup=inline_keyboard.add(InlineKeyboardButton(text="⏪Orqaga", callback_data="list_channels")).as_markup())

@dp.callback_query(lambda query: query.data.startswith("new_invite_"))
async def new_invite_link_callback_handler(callback_query: types.CallbackQuery):
    channel_id = int(callback_query.data.split('_')[-1])
    inline_keyboard = InlineKeyboardBuilder()
    inline_keyboard.add(
        InlineKeyboardButton(
            text="⏪ orqaga",
            callback_data=f"settings_{channel_id}"
        )
    )
    try:
        invite_link = await bot.create_chat_invite_link(chat_id=channel_id)
        conn = sqlite3.connect('bot.db')
        cursor = conn.cursor()
        cursor.execute("UPDATE channels SET invite_link = ? WHERE telegram_id = ?", (invite_link.invite_link, channel_id))
        conn.commit()
        conn.close()

        await callback_query.message.edit_text(f"Yangi invite link: {html.link(value='link', link=invite_link.invite_link)}",disable_web_page_preview=True,reply_markup=inline_keyboard.as_markup())
    except Exception as e:
        await callback_query.message.edit_text(f"Link yaratishda xato: {e}\nBot ruxsatlarini tekshiring", reply_markup=inline_keyboard.as_markup())


@dp.callback_query(lambda query: query.data.startswith("delete_"))
async def delete_channel_callback_handler(callback_query: types.CallbackQuery):
    channel_id = int(callback_query.data.split('_')[-1])
    inline_keyboard = InlineKeyboardBuilder()
    inline_keyboard.add(
        InlineKeyboardButton(
            text="⏪ Kanallar ro'yxatiga qaytish",
            callback_data=f"list_channels"
        )
    )
    try:
        conn = sqlite3.connect('bot.db')
        cursor = conn.cursor()
        cursor.execute("DELETE FROM channels WHERE telegram_id = ?", (channel_id,))
        conn.commit()
        conn.close()
        try:
            await bot.leave_chat(chat_id=channel_id)
        except:
            pass
        await callback_query.message.edit_text(f"Kanal o'chirildi: {channel_id}\nVa bot chatni tark etdi", reply_markup=inline_keyboard.as_markup())
    except Exception as e:
        await callback_query.message.edit_text(f"Kanalni o'chirishda xato: {e}",  reply_markup=inline_keyboard.as_markup())



@dp.message(Command('panel'), ~IsAdmin(), IsPrivate())
async def not_admin_statistic(message: types.Message):
    reaction_list = ["😁", "🤔","🤣","🤪", "🗿", "🆒","😎", "👾", "🤷‍♂", "🤷"]
    try:
        await bot.set_message_reaction(
            chat_id=message.chat.id,
            message_id=message.message_id,
            reaction=[ReactionTypeEmoji(emoji=random.choice(reaction_list))],
            is_big=False
        )
    except:
        pass
    await message.reply("Siz admin emassiz!", disable_notification=True, protect_content=True)



def check_user_in_requests(user_id: int, chat_id: int) -> bool:
    """
    Foydalanuvchining ma'lum bir kanalga qo'shilish so'rovi mavjudligini tekshirish.
    """
    with db_connection() as conn:
        cursor = conn.cursor()
        cursor.execute("""
            SELECT id FROM join_requests
            WHERE user_id = ? AND chat_id = ?
        """, (user_id, chat_id))
        result = cursor.fetchone()

    return result is not None

def get_channels():
    conn = sqlite3.connect('bot.db')
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM channels")
    channels = cursor.fetchall()
    conn.close()
    return channels



async def check_sub(user_id):
    channels = get_channels()
    for channel in channels:
        chat_id = channel[3]
        if check_user_in_requests(user_id, chat_id):
            continue
        try:
            chat_member = await bot.get_chat_member(chat_id, user_id)
            if chat_member.status not in ["creator", "administrator", "member", 'restricted']:
                return False
        except Exception as e:
            print(f"Xatolik: {e}")
            return False
    return True




async def generate_and_send_file(sorted_detailed_info, message):
    """
    Top foydalanuvchilar haqidagi ma'lumotlarni Word faylga saqlaydi va yuboradi.
    """
    try:
        document = Document()
        document.add_heading("Top Foydalanuvchilar Haqida Ma'lumot", level=1)

        for _, info in sorted_detailed_info:
            document.add_paragraph(info)

        file_name = f"top_users_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        document.save(file_name)

        await message.answer_document(
            FSInputFile(file_name),
            caption="Top foydalanuvchilar haqidagi ma'lumotlar Word fayl shaklida yuborildi."
        )

    finally:
        if os.path.exists(file_name):
            os.remove(file_name)

@dp.message(Command('top_users'), IsAdmin())
async def show_top_users(message: types.Message):
    """
    Admin uchun Top Userlarning batafsil ma'lumotlarini ko'rsatadi.
    """
    # Top taklif qilgan foydalanuvchilarni olish
    cursor.execute('''
        SELECT 
            u.id, u.full_name, u.username, u.telegram_id, u.referred_by, u.invited_users
        FROM 
            users u
        WHERE
            json_array_length(u.invited_users) > 0
        ORDER BY
            json_array_length(u.invited_users) DESC
        LIMIT 500
    ''')
    users = cursor.fetchall()

    if not users:
        await message.answer("Foydalanuvchilar ma'lumotlari topilmadi.")
        return

    response = "📊 <b>Top Foydalanuvchilar Haqida Ma'lumot:</b>\n\n"
    detailed_info = []

    for user in users:
        user_id, full_name, username, telegram_id, referred_by, invited_users = user

        # Kim tomonidan taklif qilinganini olish
        inviter_info = None
        if referred_by:
            cursor.execute('SELECT full_name, username FROM users WHERE id = ?', (referred_by,))
            inviter_info = cursor.fetchone()

        # Taklif qilgan foydalanuvchilarni aniqlash
        invited_users_list = json.loads(invited_users) if invited_users else []
        
        valid_invited_users = []
        for invited_id in invited_users_list:
            if await check_sub(invited_id):
                valid_invited_users.append(invited_id)

        # Har bir user haqida ma'lumotni tayyorlash
        user_info = f"👤 {full_name} (@" \
                    f"{username if username else 'username yo‘q'})\n"
        user_info += f"🆔 Telegram ID: {telegram_id}\n"
        user_info += f"👥 Taklif qilgan foydalanuvchilar soni: {len(invited_users_list)}\n"
        user_info += f"✅ Obuna bo'lgan takliflar soni: {len(valid_invited_users)}\n"

        if inviter_info:
            user_info += f"🔗 Taklif qilgan: {inviter_info[0]} (@{inviter_info[1]})\n"

        user_info += "\n"
        detailed_info.append((len(valid_invited_users), user_info))  # valid_invited_users soni bo'yicha sortlash

    # Eng ko'p taklif qilgan foydalanuvchilarni yuqoriga joylashtirish
    sorted_detailed_info = sorted(detailed_info, reverse=True, key=lambda x: x[0])

    # Javobni cheklash va kerak bo'lsa Word faylga saqlash
    if sum(len(info[1]) for info in sorted_detailed_info) > 400:
        await generate_and_send_file(sorted_detailed_info, message)

    else:
        # Faylni yuklab olish uchun tugmani qo‘shish
        file_id = str(uuid.uuid4())  # Faylni noyob qilish uchun ID yaratish

        response += "\n".join(info for _, info in sorted_detailed_info)
        await message.answer(response, 
        reply_markup = InlineKeyboardMarkup(
            inline_keyboard=[
                [InlineKeyboardButton(text="Faylni shaklida yuklab olish", callback_data=f"download_file_{file_id}")]
            ]
        ),)



@dp.callback_query(lambda c: c.data and c.data.startswith('download_file_'))
async def handle_download_file(callback_query: types.CallbackQuery):
    """
    Callback handler for downloading the file when the user clicks the inline button.
    """
    file_id = callback_query.data.split('_')[2]  # Faylni ID dan olish
    # Faylni yaratib, yuborish
    # Bu yerda `generate_and_send_file` funksiyasini chaqirishingiz mumkin.
    # Agar fayl hozirda mavjud bo'lsa, uni yuborish.
    
    # Barcha kerakli ma'lumotlar va faylni yuborish
    cursor.execute('''
        SELECT 
            u.id, u.full_name, u.username, u.telegram_id, u.referred_by, u.invited_users
        FROM 
            users u
        WHERE
            json_array_length(u.invited_users) > 0
        ORDER BY
            json_array_length(u.invited_users) DESC
        LIMIT 500
    ''')
    users = cursor.fetchall()
    
    if not users:
        await callback_query.answer("Foydalanuvchilar ma'lumotlari topilmadi.")
        return

    detailed_info = []
    for user in users:
        user_id, full_name, username, telegram_id, referred_by, invited_users = user
        inviter_info = None
        if referred_by:
            cursor.execute('SELECT full_name, username FROM users WHERE id = ?', (referred_by,))
            inviter_info = cursor.fetchone()

        invited_users_list = json.loads(invited_users) if invited_users else []
        valid_invited_users = []
        for invited_id in invited_users_list:
            if await check_sub(invited_id):
                valid_invited_users.append(invited_id)

        user_info = f"👤 {full_name} (@" \
                    f"{username if username else 'username yo‘q'})\n"
        user_info += f"🆔 Telegram ID: {telegram_id}\n"
        user_info += f"👥 Taklif qilgan foydalanuvchilar soni: {len(invited_users_list)}\n"
        user_info += f"✅ Obuna bo'lgan takliflar soni: {len(valid_invited_users)}\n"

        if inviter_info:
            user_info += f"🔗 Taklif qilgan: {inviter_info[0]} (@{inviter_info[1]})\n"

        user_info += "\n"
        detailed_info.append((len(valid_invited_users), user_info))

    sorted_detailed_info = sorted(detailed_info, reverse=True, key=lambda x: x[0])
    await generate_and_send_file(sorted_detailed_info, callback_query.message)
    await callback_query.answer()