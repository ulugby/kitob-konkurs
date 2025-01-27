from aiogram.filters import CommandStart, CommandObject, Command
from loader import dp,bot
from aiogram import types,html
from aiogram.utils.keyboard import InlineKeyboardBuilder
import sqlite3
from datetime import datetime
from aiogram.types import  InlineKeyboardButton,InlineKeyboardMarkup, ReplyKeyboardRemove
from aiogram.types.reaction_type_emoji import ReactionTypeEmoji
import random
from data.config import ADMINS, USERS_CHANNEL
import json,os

from utils.misc.checksub import joinchat
from docx import Document

from aiogram.types import FSInputFile

from urllib.parse import quote


DATABASE_FILE = "bot.db"

conn = sqlite3.connect('bot.db')
cursor = conn.cursor()

def db_connection():
    return sqlite3.connect(DATABASE_FILE)


def check_user_in_requests(user_id: int, chat_id: int) -> bool:
    """
    Foydalanuvchining ma'lum bir kanalga qo'shilish so'rovi mavjudligini tekshirish.
    """
    with db_connection() as conn:
        cursor = conn.cursor()
        cursor.execute("""
            SELECT id FROM join_requests
            WHERE user_id = ? AND chat_id = ?
        """, (user_id, chat_id))
        result = cursor.fetchone()

    return result is not None

def get_channels():
    conn = sqlite3.connect('bot.db')
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM channels")
    channels = cursor.fetchall()
    conn.close()
    return channels



async def check_sub(user_id):
    channels = get_channels()
    for channel in channels:
        chat_id = channel[3]
        if check_user_in_requests(user_id, chat_id):
            continue
        try:
            chat_member = await bot.get_chat_member(chat_id, user_id)
            if chat_member.status not in ["creator", "administrator", "member", 'restricted']:
                return False
        except Exception as e:
            print(f"Xatolik: {e}")
            return False
    return True





cursor.execute('''
CREATE TABLE IF NOT EXISTS users (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    username VARCHAR(30) NULL,
    full_name TEXT,
    telegram_id BIGINTEGER UNIQUE,
    lang_code VARCHAR(10) NULL,
    registration_date TEXT,
    referred_by INTEGER,
    referral_code VARCHAR(20) NULL UNIQUE,
    invited_users TEXT DEFAULT '[]',
    FOREIGN KEY (referred_by) REFERENCES users (id)
)
''')
conn.commit()

def generate_referral_code():
    """Tasodifiy referal kod yaratish."""
    import uuid
    return f"r_{uuid.uuid4().hex[:8]}"


def get_user_by_referral_code(referral_code):
    cursor.execute('''
        SELECT * FROM users WHERE referral_code = ?
    ''', (referral_code,))
    return cursor.fetchone()

def add_user(telegram_id, username, full_name, registration_date, referred_by=None):
    referral_code = generate_referral_code()
    invited_users = json.dumps([])  # Taklif qilingan do'stlar ro'yxati
    cursor.execute('''
        INSERT INTO users (telegram_id, username, full_name, registration_date, referred_by, invited_users, referral_code)
        VALUES (?, ?, ?, ?, ?, ?, ?)
    ''', (telegram_id, username, full_name, registration_date, referred_by, invited_users, referral_code))
    conn.commit()

    # Agar referred_by mavjud bo'lsa, taklif qilgan foydalanuvchining invited_users maydonini yangilaymiz
    if referred_by:
        cursor.execute('SELECT invited_users FROM users WHERE id = ?', (referred_by,))
        result = cursor.fetchone()
        if result:
            invited_list = json.loads(result[0])  # Taklif qilingan foydalanuvchilar ro'yxati
            invited_list.append(telegram_id)  # Yangi foydalanuvchini qo'shamiz
            updated_invited_users = json.dumps(invited_list)
            cursor.execute('UPDATE users SET invited_users = ? WHERE id = ?', (updated_invited_users, referred_by))
            conn.commit()


# Foydalanuvchini tekshirish
async def is_user_registered(telegram_id):
    cursor.execute('SELECT telegram_id FROM users WHERE telegram_id = ?', (telegram_id,))
    return cursor.fetchone() is not None


def get_user_referral_code(telegram_id):
    cursor.execute('SELECT referral_code FROM users WHERE telegram_id = ?', (telegram_id,))
    result = cursor.fetchone()
    if result:
        return result[0]
    else:
        return None


def get_top_referrers(limit=50):
    cursor.execute('''
        SELECT username, full_name, referral_code, json_array_length(invited_users) AS invited_count
        FROM users
        WHERE json_array_length(invited_users) > 0
        ORDER BY invited_count DESC
        LIMIT ?
    ''', (limit,))
    return cursor.fetchall()


# /start komandasi
@dp.message(CommandStart())
async def start_bot(message: types.Message, command: CommandObject):
    telegram_id = message.from_user.id
    max_length = 20
    full_name = message.from_user.full_name
    full_name = (full_name[:max_length] + '...') if len(full_name) > max_length else full_name
    username = message.from_user.username
    registration_date = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    args = command.args  # Referal kodni o'qish
    is_premium = message.from_user.is_premium



    welcome_text = (
        f"👋 Assalamu alaykum, {html.bold(full_name)}!\n\n"
        "Sizni botimizda ko'rganimizdan xursandmiz! 🤖\n\n"
        "📚 Ushbu bot orqali siz:\n"
        "   • Do'stlaringizni taklif qilib, <b>'Muqaddima'</b> kitobini sovg'a sifatida yutib olishingiz mumkin. 🎁\n"
        "   • Eng faol taklif qiluvchilar orasida o'z o'rningizni topishingiz mumkin. 🏆\n\n"
        "🛠 Botning asosiy buyruqlari:\n"
        "/start - Botni qayta boshlash va ma'lumot olish.\n"
        "/referal - Taklif qilgan do'stlaringiz haqida ma'lumot.\n"
        "/topreferals - Eng faol taklif qiluvchilar ro'yxati.\n\n"
        "👉 Boshlash uchun, do'stlaringizni taklif qiling va sovrinlarni qo'lga kiriting! 🎉\n"
        f"📢 Kanalimizga obuna bo'ling: {html.link(value='Bizning Kanal', link='https://t.me/kariymulloh')}\n"
        "❓Bot Savollaringiz bo'lsa, qo'llab-quvvatlash: @web_icon\n"
    )

    if not await is_user_registered(telegram_id):  # Agar foydalanuvchi ro'yxatda bo'lmasa
        referred_by = None
        if args:  # Agar referal kod bo'lsa
            referred_user = get_user_by_referral_code(args)
            if referred_user:
                referred_by = referred_user[0]  # Taklif qilgan foydalanuvchining ID-si

        # Foydalanuvchini qo'shamiz
        add_user(telegram_id, username, full_name, registration_date, referred_by)

        try:
            await bot.send_message(
                chat_id=USERS_CHANNEL,
                text=(
                    f"New 👤: {full_name}\n"
                    f"Username📩: {html.code(username)}\n"
                    f"Telegram 🆔: {html.code(str(telegram_id))}\n"
                    f"Reg 📆: {registration_date}\n"
                    f"Premium🤑: {is_premium}"
                ),
                reply_markup=InlineKeyboardMarkup(
                    inline_keyboard=[[InlineKeyboardButton(text="Profile", url=f"tg://user?id={telegram_id}")]]
                ),
            )
        except:
            for admin_id in ADMINS:
                try:
                    await bot.send_message(
                        chat_id=admin_id,
                        text=(
                            f"New 👤: {full_name}\n"
                            f"Username📩: {html.code(username)}\n"
                            f"Telegram 🆔: {html.code(str(telegram_id))}\n"
                            f"Reg 📆: {registration_date}\n"
                            f"Premium🤑: {is_premium}"
                        ),
                        reply_markup=InlineKeyboardMarkup(
                            inline_keyboard=[[InlineKeyboardButton(text="Profile", url=f"tg://user?id={telegram_id}")]]
                        ),
                    )
                except:
                    pass
        reaction_list = ["❤", "🔥", "🥰","🎉", "🤩","🕊","😍", "❤‍🔥", "🌚","⚡","🤗"]
        try:
            await bot.set_message_reaction(
                chat_id=message.chat.id,
                message_id=message.message_id,
                reaction=[ReactionTypeEmoji(emoji=random.choice(reaction_list))],
                is_big=False
            )
        except:
            pass

        await message.answer(welcome_text, disable_web_page_preview=True, reply_markup=ReplyKeyboardRemove())
    is_subscribed = await joinchat(message.from_user.id)
    if not is_subscribed:
        return
    
    bot_username = (await bot.me()).username
    referral_link = f"https://t.me/{bot_username}?start={get_user_referral_code(telegram_id)}\n\n"

    post_text = "📚 Muqqadima kitobini yutib olish imkoniyati! 🎉\n\nDo'stlaringizni taklif qilib, ular bilan birga yutish imkoniyatini oshiring!"

    
    user_text =  f"📚 <b>Muqqadima kitobini yutib olish imkoniyati!'</b> 🎉\n\n" \
                f"Do'stlaringizni taklif qiling, yutish imkoniyatini oshiring! 🎁\n\n" \
                f'👇 <b>"Quyidagi havolani do\'stlaringizga yuboring!" </b> 👇'

    encoded_referral_link = quote(referral_link, safe="")
    encoded_post_text = quote(post_text, safe="")

    # Xabarni yuborish va ulashish tugmasini qo'shish
    await message.answer(
        user_text,
        reply_markup=InlineKeyboardMarkup(
            inline_keyboard=[
                [InlineKeyboardButton(text="Do'stlaringizni Taklif Qilish", url=f"https://t.me/share/url?url={encoded_referral_link}&text={encoded_post_text}")]
            ]
        ),
    )
@dp.message(Command('referal'))
async def show_referral_info(message: types.Message):
    telegram_id = message.from_user.id

    cursor.execute('SELECT id, full_name, referred_by, invited_users FROM users WHERE telegram_id = ?', (telegram_id,))
    user = cursor.fetchone()

    if not user:
        return

    user_id, full_name, referred_by, invited_users = user

    invited_users_list = json.loads(invited_users) if invited_users else []
    valid_invited_users = [u for u in invited_users_list if await check_sub(u)]



    # Kim taklif qilganini topish
    inviter_info = None
    if referred_by:
        cursor.execute('SELECT full_name, username FROM users WHERE id = ?', (referred_by,))
        inviter_info = cursor.fetchone()

    # Taklif qilgan foydalanuvchilarni olish
    invited_info = []
    if invited_users_list:
        cursor.execute(
            f"SELECT full_name, username, telegram_id FROM users WHERE telegram_id IN ({','.join('?' * len(invited_users_list))})",
            invited_users_list,
        )
        invited_info = cursor.fetchall()

    # Javobni tayyorlash
    response = f"👤 Siz haqingizda ma'lumot:\n"
    response += f"🆔 Telegram ID: {telegram_id}\n"
    response += f"📅 Taklif qilganlar soni: {len(invited_users_list)}\n"
    response += f"📅 Kanalga obuna bo'lgan taklif qilganlar soni: {len(valid_invited_users)}\n"

    

    if inviter_info:
        response += f"\n🧑‍💼 Sizni taklif qilgan: {inviter_info[0]} (@{inviter_info[1]})\n"

    if invited_info:
        response += "\n👥 Siz taklif qilgan foydalanuvchilar:\n"
        for idx, (full_name, username, t_id) in enumerate(invited_info, start=1):
            username_result = f"(@{username})" if username else ""
            response += f"{idx}. {full_name} {username_result} — 🆔 {t_id}\n"

    if len(response) > 1020:
        try:
            # Word fayl yaratish
            document = Document()
            today = datetime.now().strftime('%Y-%m-%d')
            document.add_heading(f"Referal Ma'lumotlar ({today})", level=1)
            document.add_paragraph(f"👤 Ism: {full_name}")
            document.add_paragraph(f"🆔 Telegram ID: {telegram_id}")
            document.add_paragraph(f"📅 Taklif qilganlar soni: {len(invited_users_list)}")
            document.add_paragraph(f"📅 Kanalga obuna bo'lgan taklif qilganlar soni: {len(valid_invited_users)}")

            if inviter_info:
                document.add_paragraph(f"🧑‍💼 Sizni taklif qilgan: {inviter_info[0]} (@{inviter_info[1]})")

            if invited_info:
                document.add_heading("👥 Kanalga obuna bo'lgan Taklif qilingan foydalanuvchilar:", level=2)
                for idx, (inv_full_name, inv_username, inv_telegram_id) in enumerate(valid_invited_users, start=1):
                    document.add_paragraph(
                        f"{inv_full_name} (@{inv_username}) — 🆔 {inv_telegram_id}",
                        style='List Number'
                    )

            # Fayl nomini yaratish
            file_name = f"referal_info_{telegram_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

        finally:
            if os.path.exists(file_name):
                os.remove(file_name)

            await message.answer_document(FSInputFile(file_name), caption="Sizning referal ma'lumotlaringiz.")

            # Faylni o'chirish
            os.remove(file_name)
    else:
        # Javobni xabar sifatida yuborish
        await message.answer(response)


@dp.message(Command('topreferals'))
async def show_top_referrers(message: types.Message):
    top_referrers = get_top_referrers()

    if not top_referrers:
        await message.answer("Hozircha hech kim do'st taklif qilmagan. 😊")
        return

    # Top foydalanuvchilarni formatlash
    result_text = "<b>Eng ko'p do'st taklif qilgan foydalanuvchilar:</b>\n\n"
    for i, user in enumerate(top_referrers, start=1):
        username = f"@{user[0]}" if user[0] else "Noma'lum"
        full_name = user[1] if user[1] else "Noma'lum"
        invited_count = user[3]
        result_text += (
            f"{i}. {full_name} ({username})\n"
            f"   Taklif qilganlar: {invited_count}\n\n"
        )

    # Agar xabar uzunligi 1024 dan oshsa, Word fayl yaratish
    if len(result_text) > 1020:
        try:
            document = Document()
            document.add_heading("Eng ko'p do'st taklif qilgan foydalanuvchilar", level=1)

            for i, user in enumerate(top_referrers, start=1):
                username = f"@{user[0]}" if user[0] else "Noma'lum"
                full_name = user[1] if user[1] else "Noma'lum"
                invited_count = user[3]
                document.add_paragraph(
                    f"{i}. {full_name} ({username})\n"
                    f"   Taklif qilganlar: {invited_count}\n"
                )

            # Word faylni saqlash
            file_name = f"top_referrals_{message.from_user.id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            document.save(file_name)

            # Faylni foydalanuvchiga yuborish
            await message.answer_document(
                FSInputFile(file_name),
                caption="Eng ko'p do'st taklif qilgan foydalanuvchilar ro'yxati Word fayl ko'rinishida."
            )

        finally:
            if os.path.exists(file_name):
                os.remove(file_name)
    else:
        # Xabarni oddiy matn sifatida yuborish
        await message.answer(result_text)